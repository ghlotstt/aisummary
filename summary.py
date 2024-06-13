'''
# summary.py

import os
import requests
from dotenv import load_dotenv

# Cargar las variables de entorno
load_dotenv()

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
OPENAI_API_URL = "https://api.openai.com/v1/chat/completions"

def summarize_text(text):
    """Generar un resumen utilizando GPT-3.5 a través de una solicitud HTTP."""
    headers = {
        "Authorization": f"Bearer {OPENAI_API_KEY}",
        "Content-Type": "application/json"
    }
    
    prompt = f"Please summarize the following conversation:\n\n{text}"
    
    data = {
        "model": "gpt-3.5-turbo",
        "messages": [
            {"role": "system", "content": "You are an assistant who summarizes conversations."},
            {"role": "user", "content": prompt}
        ],
        "max_tokens": 300,
        "temperature": 0.5
    }
    
    try:
        response = requests.post(OPENAI_API_URL, headers=headers, json=data)
        response.raise_for_status()
        return response.json()["choices"][0]["message"]["content"].strip()
    except requests.exceptions.RequestException as e:
        print(f"Error summarizing text: {e}")
        return "Error summarizing the conversation."
'''


# summary.py
import os
import requests
from dotenv import load_dotenv
from docx import Document

# Cargar las variables de entorno
load_dotenv()

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
OPENAI_API_URL = "https://api.openai.com/v1/chat/completions"

PROMPTS = {
    "sermon": {
        "role": "system",
        "content": ("You are an assistant specialized in summarizing sermons given in Christian churches. "
                    "Your summaries should: "
                    "1. Capture the core messages, including spiritual teachings, biblical references, y practical applications para la vida cotidiana. "
                    "2. Identify the themes of the sermon (e.g., faith, love, repentance). "
                    "3. Provide a summary in the following format:\n"
                    "   - **Title**: [Title of the sermon]\n"
                    "   - **Main Themes**: [Key themes identified]\n"
                    "   - **Biblical References**: [Relevant biblical passages or verses]\n"
                    "   - **Summary**: Write a cohesive narrative that integrates the core message, themes, and biblical references directly, without phrases like 'The sermon emphasizes'. Focus on conveying the message and teachings.\n"
                    "   - **Powerful Keywords**: A list of up to 10 keywords that capture the sermon’s essence.\n"
                    "   - **Powerful Phrases**: Up to 5 powerful and inspiring frases que resuenen con los oyentes.\n"
                    "\n"
                    "Here is the sermon text:\n\n"
                    "{text}\n\n"
                    "Provide the summary and keywords/phrases in this format:\n"
                    "**Title**: [Title of the sermon]\n"
                    "**Main Themes**: [List of main themes]\n"
                    "**Biblical References**: [Relevant biblical passages or verses]\n"
                    "**Summary**:\n"
                    "**Powerful Keywords**:\n"
                    "**Powerful Phrases**:")
    },
    "lecture": {
        "role": "system",
        "content": ("You are an assistant specialized in summarizing TED talks. "
                                     "Your summaries should: "
                                     "1. Capture the core messages, including key insights and important messages. "
                                     "2. Identify the themes of the talk (e.g., innovation, science, culture). "
                                     "3. Provide a summary in the following format:\n"
                                     "   - **Title**: [Title of the talk]\n"
                                     "   - **Main Themes**: [Key themes identified]\n"
                                     "   - **Key Points**: [Important points or messages]\n"
                                     "   - **Summary**: [Detailed summary capturing the core message]\n"
                                     "4. Generate a list of powerful keywords and phrases that will inspire and engage the audience, helping them remember the talk more clearly.\n"
                                     "   - **Powerful Keywords**: A list of up to 10 keywords that capture the talk’s essence.\n"
                                     "   - **Powerful Phrases**: Up to 5 powerful and inspiring phrases that resonate with the audience.\n"
                                     "\n"
                                     "Here is the talk text:\n\n"
                                     "{text}\n\n"
                                     "Provide the summary and keywords/phrases in this format:\n"
                                     "**Title**: [Title of the talk]\n"
                                     "**Main Themes**: [List of main themes]\n"
                                     "**Key Points**: [Relevant points or messages]\n"
                                     "**Summary**:\n"
                                     "**Powerful Keywords**:\n"
                                     "**Powerful Phrases**:")
    },
    "conference": {
        "role": "system",
        "content": ("You are an assistant specialized in summarizing TED talks. "
                                     "Your summaries should: "
                                     "1. Capture the core messages, including key insights and important messages. "
                                     "2. Identify the themes of the talk (e.g., innovation, science, culture). "
                                     "3. Provide a summary in the following format:\n"
                                     "   - **Title**: [Title of the talk]\n"
                                     "   - **Main Themes**: [Key themes identified]\n"
                                     "   - **Key Points**: [Important points or messages]\n"
                                     "   - **Summary**: [Detailed summary capturing the core message]\n"
                                     "4. Generate a list of powerful keywords and phrases that will inspire and engage the audience, helping them remember the talk more clearly.\n"
                                     "   - **Powerful Keywords**: A list of up to 10 keywords that capture the talk’s essence.\n"
                                     "   - **Powerful Phrases**: Up to 5 powerful and inspiring phrases that resonate with the audience.\n"
                                     "\n"
                                     "Here is the talk text:\n\n"
                                     "{text}\n\n"
                                     "Provide the summary and keywords/phrases in this format:\n"
                                     "**Title**: [Title of the talk]\n"
                                     "**Main Themes**: [List of main themes]\n"
                                     "**Key Points**: [Relevant points or messages]\n"
                                     "**Summary**:\n"
                                     "**Powerful Keywords**:\n"
                                     "**Powerful Phrases**:")
    }
}

def summarize_text(text, summary_type="sermon"):
    """Generar un resumen utilizando GPT-4o a través de una solicitud HTTP."""
    headers = {
        "Authorization": f"Bearer {OPENAI_API_KEY}",
        "Content-Type": "application/json"
    }
    
    prompt_message = PROMPTS.get(summary_type, PROMPTS["sermon"])
    
    data = {
        "model": "gpt-4o",
        "messages": [
            prompt_message,
            {"role": "user", "content": prompt_message['content'].format(text=text)}
        ],
        "max_tokens": 1000,
        "temperature": 0.5
    }
    
    try:
        response = requests.post(OPENAI_API_URL, headers=headers, json=data)
        response.raise_for_status()
        return response.json()["choices"][0]["message"]["content"].strip()
    except requests.exceptions.RequestException as e:
        print(f"Error summarizing text: {e}")
        return f"Error summarizing the {summary_type}."

#==== Save Summary Function ===================================
    
def save_summary_to_docx(summary, filename='summary.docx'):
    """Guardar el resumen en un documento .docx."""
    doc = Document()
    doc.add_heading('Summary', 0)
    doc.add_paragraph(summary)
    doc.save(filename)
    print(f'Summary saved to {os.path.abspath(filename)}')
'''

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def save_summary_to_docx(summary, filename='summary.docx'):
    doc = Document()
    doc.add_heading('Summary', level=1)

    # Title
    title = summary.get('title', 'No Title Provided')
    doc.add_heading(title, level=2)

    # Main Themes
    themes = summary.get('themes', 'No Themes Provided')
    doc.add_heading('Main Themes:', level=3)
    doc.add_paragraph(themes)

    # Key Points
    points = summary.get('key_points', 'No Key Points Provided')
    doc.add_heading('Key Points:', level=3)
    for point in points:
        doc.add_paragraph(point, style='ListBullet')

    # Detailed Summary
    detailed_summary = summary.get('detailed_summary', 'No Detailed Summary Provided')
    doc.add_heading('Detailed Summary:', level=3)
    doc.add_paragraph(detailed_summary)

    # Keywords and Phrases
    doc.add_heading('Powerful Keywords:', level=3)
    keywords = summary.get('keywords', [])
    doc.add_paragraph(', '.join(keywords))

    doc.add_heading('Powerful Phrases:', level=3)
    phrases = summary.get('phrases', [])
    for phrase in phrases:
        doc.add_paragraph(phrase, style='Quote')

    doc.save(filename)
    print(f'Summary saved to {filename}')
'''